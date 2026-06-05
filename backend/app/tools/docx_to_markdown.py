"""
File Handling Tools — DOCX to Markdown Converter.

docx_to_markdown  → convert Microsoft Word documents to clean Markdown
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

from docx import Document
from pydantic import Field

from app.tools._file_utils import resolve_input
from app.tools.base import BaseTool, ToolInput, ToolMetadata, ToolResult, register_tool

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _extract_paragraph_text(para) -> str:
    """Extract text from a paragraph, preserving inline formatting as Markdown."""
    text_parts: list[str] = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            t = f"***{t}***"
        elif run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        if run.font.strike:
            t = f"~~{t}~~"
        if run.underline:
            t = f"<u>{t}</u>"
        text_parts.append(t)
    return "".join(text_parts)


_BULLET_RE = re.compile(r"^\s*[\u2022\u2023\u25E6\u2043\u2219\-\*]\s*")
_NUMBERED_RE = re.compile(r"^\s*\d+[\.\)]\s*")


def _detect_list_prefix(text: str) -> str:
    """Detect list type and return the Markdown prefix string (e.g. "- " or "1. ").

    Returns "" if no list prefix is found.
    """
    m_b = _BULLET_RE.match(text)
    if m_b:
        return "- "
    m_n = _NUMBERED_RE.match(text)
    if m_n:
        num_str = m_n.group().strip().rstrip(".)")
        return f"{num_str}. "
    return ""


def _format_paragraph_markdown(para) -> str:
    """Convert a docx paragraph to Markdown, handling headings, lists, and inline formatting."""
    style = para.style.name.lower() if para.style and para.style.name else ""

    # Headings
    if style.startswith("heading"):
        try:
            level = int(style.split()[-1])
        except ValueError:
            level = 1
        text = _extract_paragraph_text(para).strip()
        if text:
            return f"{'#' * level} {text}\n"
        return ""

    # Detect list prefix from the formatted (Markdown) text directly
    formatted = _extract_paragraph_text(para)
    prefix = _detect_list_prefix(formatted)
    if prefix:
        # Strip the matched prefix characters from formatted text
        # Match leading whitespace + bullet/number + whitespace
        m = re.match(r"^\s*(?:[-\u2022\u2023\u25E6\u2043\u2219\*]\s*|\d+[\.\)]\s*)", formatted)
        if m:
            formatted = formatted[m.end():]
        return f"{prefix}{formatted.strip()}\n"

    # Normal paragraph
    formatted = _extract_paragraph_text(para)
    if formatted.strip():
        return f"{formatted}\n"
    return "\n"


def _table_to_markdown(table) -> str:
    """Render a python-docx table as GitHub-flavored Markdown."""
    if not table.rows:
        return ""

    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    lines: list[str] = []
    # Header
    lines.append("| " + " | ".join(rows[0]) + " |")
    # Separator
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Body
    for row in rows[1:]:
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")

    lines.append("")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# docx_to_markdown
# ---------------------------------------------------------------------------


class DocxToMarkdownInput(ToolInput):
    data: str | None = Field(
        None,
        description="Base64-encoded DOCX content (optional if 'url' is provided)",
    )
    url: str | None = Field(
        None,
        description="URL to fetch the DOCX from (optional if 'data' is provided)",
    )
    include_tables: bool = Field(
        True,
        description="Render tables as GitHub-flavored Markdown tables",
    )
    include_images: bool = Field(
        False,
        description="Extract embedded images as base64 data URIs (increases output size)",
    )
    preserve_track_changes: bool = Field(
        False,
        description="Include revision tracking markup in output",
    )


class DocxToMarkdownTool(BaseTool):
    def __init__(self):
        metadata = ToolMetadata(
            tool_id="docx_to_markdown",
            name="DOCX to Markdown",
            description="Convert Microsoft Word (.docx) documents to clean Markdown for LLM consumption",
            category="file-handling",
            input_schema=DocxToMarkdownInput.schema_extra(),
            tags=["docx", "markdown", "convert", "word", "file-handling"],
            requires_auth=True,
        )
        super().__init__(metadata=metadata)

    async def execute(self, input_data: dict) -> ToolResult:
        try:
            validated = DocxToMarkdownInput(**input_data)
        except Exception as e:
            return ToolResult.error_result(
                tool_id=self.tool_id, error=f"Invalid input: {e}"
            )

        try:
            docx_bytes = await resolve_input(
                validated.data, validated.url, label="DOCX"
            )
        except ValueError as e:
            return ToolResult.error_result(tool_id=self.tool_id, error=str(e))
        except Exception as e:
            return ToolResult.error_result(
                tool_id=self.tool_id, error=f"Failed to read DOCX: {e}"
            )

        try:
            doc = Document(io.BytesIO(docx_bytes))
            md_lines: list[str] = []
            image_count = 0

            # Document title from properties
            if doc.core_properties.title:
                md_lines.append(f"# {doc.core_properties.title}\n")

            # Iterate body using counters (O(n) — avoids repeated scans)
            p_idx = 0
            tbl_idx = 0
            for element in doc.element.body:
                tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

                if tag == "p":
                    if p_idx < len(doc.paragraphs):
                        para = doc.paragraphs[p_idx]
                        p_idx += 1
                        md_lines.append(_format_paragraph_markdown(para))

                elif tag == "tbl" and validated.include_tables and tbl_idx < len(doc.tables):
                    table = doc.tables[tbl_idx]
                    tbl_idx += 1
                    md_lines.append(_table_to_markdown(table))

            markdown = "\n".join(md_lines)

            # Collapse excessive blank lines
            markdown = re.sub(r"\n{3,}", "\n\n", markdown)

            result: dict[str, Any] = {
                "markdown": markdown,
                "character_count": len(markdown),
                "word_count": len(markdown.split()),
                "line_count": len(md_lines),
                "paragraph_count": len(doc.paragraphs),
            }

            if validated.include_images and image_count:
                result["embedded_images"] = image_count

            # Document metadata
            try:
                props = doc.core_properties
                result["metadata"] = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                    "last_modified_by": props.last_modified_by or "",
                    "revision": props.revision or 0,
                }
            except Exception:
                result["metadata"] = {}

            return ToolResult.success_result(tool_id=self.tool_id, result=result)

        except Exception as e:
            logger.exception("docx_to_markdown failed")
            return ToolResult.error_result(tool_id=self.tool_id, error=str(e))


# ---------------------------------------------------------------------------
# register
# ---------------------------------------------------------------------------

register_tool(DocxToMarkdownTool())
